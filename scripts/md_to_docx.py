"""Convert SARa Plan.md (report markdown) to a styled .docx.
Handles: headings, tables, bold/italic/inline-code, links (text only),
blockquotes, bullet/numbered lists, horizontal rules.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1]
DST = sys.argv[2]

INLINE_RE = re.compile(
    r"(\*\*.+?\*\*)"          # bold
    r"|(`[^`]+`)"             # inline code
    r"|(\[[^\]]+\]\([^)]+\))" # link
    r"|(\*[^*]+\*)"           # italic
)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text, base_italic=False):
    """Parse inline markdown and append formatted runs to paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.italic = base_italic
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
            r.italic = base_italic
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif tok.startswith("["):
            lm = LINK_RE.match(tok)
            label, url = lm.group(1), lm.group(2)
            r = paragraph.add_run(label)
            r.italic = base_italic
        elif tok.startswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.italic = base_italic


def is_table_sep(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells and all(set(c) <= set("-: ") and "-" in c for c in cells)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # Heading
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, hm.group(2))
            i += 1
            continue

        # Table: current line has | and next line is a separator
        if stripped.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Table Grid"
            for j, cell_text in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], cell_text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j in range(ncol):
                    txt = row[j] if j < len(row) else ""
                    cells[j].paragraphs[0].text = ""
                    add_runs(cells[j].paragraphs[0], txt)
            doc.add_paragraph()
            continue

        # Blockquote
        if stripped.startswith(">"):
            content = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, content.strip(), base_italic=True)
            i += 1
            continue

        # Bullet list
        bm = re.match(r"^[-*]\s+(.*)$", stripped)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bm.group(1))
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^\d+\.\s+(.*)$", stripped)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, nm.group(1))
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
